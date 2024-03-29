import torch
import streamlit as st
from transformers import AutoTokenizer, pipeline
from transformers import PegasusForConditionalGeneration, PegasusTokenizer
import os
import json
import docx
from zipfile import ZipFile
import shutil

HEADINGS = ["TITLE", "ABSTRACT", "RATIONALE AND BACKGROUND", "POPULATION", "RESEARCH QUESTION AND OBJECTIVES", "STUDY DESIGN", "STUDY SIZE", "DATA ANALYSIS"]
ABS_HEADINGS =  ["TITLE", "RATIONALE AND BACKGROUND",  "RESEARCH QUESTION AND OBJECTIVES", "STUDY DESIGN", "POPULATION", "DATA SOURCE", "STUDY SIZE", "DATA ANALYSIS"]

'''
Extracts the sections in the given word document of the protocol
Requirements: must be a docx, must have sections in 'header 1' formatting, 
    section headers must be bolded (?), no bullet points for best parsing results
Returns: 
    the dictionary organized into each section as a header, then abstract: and section: per each area
'''
def process_full_word_doc(file_path):
    big_dict = {}
    small_dict = {}
    headings_sections = extract_headings_sections(file_path)
    # if this is for "doc processing" for dataset, process_abstract returns the sections
    # if the abstract doesn't exist, will return None
    abstract_sections = process_abstract(file_path)
    for heading, section in headings_sections:
        for wanted in ABS_HEADINGS:
            wanted = wanted.upper()
            heading = heading.upper()
            if wanted in heading:
                small_dict['section'] = section
                if abstract_sections:
                    small_dict['abstract'] = abstract_sections[wanted.lower()]
                big_dict[wanted] = small_dict.copy()
    return big_dict
    

''' 
Run pegasus on a json file, where the sections of the file
are the sections of the protocol. This produces an abstractive
summary using pegasus of each section of the protocol. 

Input: Json file with sections of protocol OR dict of values
Output: Json file with protocol sections as headers and 
    text is the summaries generated by pegasus
'''
def run_pegasus(file, tokenizer, model):
    dict = {}
    try:
        with open(file, 'r') as f:
            dict = json.load(f)
    except TypeError: 
        dict = file
    output_dict = {}
    for header, section_abstract in dict.items():
        for key, value in section_abstract.items():
            inputs = tokenizer(value, padding='max_length', truncation=True, return_tensors='pt', max_length = 500)
            if len(value) < 500:
                summary_ids = model.generate(inputs['input_ids'], max_new_tokens=300)
            else:
                summary_ids = model.generate(inputs['input_ids'], max_new_tokens=300, encoder_repetition_penalty=1.5)
            output = tokenizer.batch_decode(summary_ids)
            output_dict[header] = output
    return output_dict

def run_llama2(file, tokenizer, model):
    dict = {}
    try:
        with open(file, 'r') as f:
            dict = json.load(f)
    except TypeError: 
        dict = file
    output_dict = {}
    for header, section_abstract in dict.items():
        for key, value in section_abstract.items():
            inputs = tokenizer(value, padding='max_length', truncation=True, return_tensors='pt', max_length = 500)
            summary_ids = model.generate(inputs['input_ids'], max_new_tokens=300, encoder_repetition_penalty=1.5)
            output = tokenizer.batch_decode(summary_ids)
            output_dict[header] = output
    return output_dict


'''
Helper functions
'''
def is_heading(paragraph):
    return paragraph.style.name.startswith("Heading 1") or paragraph.style.name.startswith("Heading 2")

'''
Returns a tuple (name of heading, heading contents) in a list
'''
def extract_headings_sections(doc_path):
    headings_sections = []
    doc = docx.Document(doc_path)
    current_heading = None
    current_section = []
    for paragraph in doc.paragraphs:
        if is_heading(paragraph):
            # if the current para is a heading, save previous section
            if current_heading and current_heading.upper() != "ABSTRACT":
                headings_sections.append((current_heading.strip(), '\n'.join(current_section)))

            # start new section for next heading
            current_heading = paragraph.text
            current_section = []

        else:
            # Collect text for current section
            current_section.append(paragraph.text.strip())
    if current_heading and current_heading.upper() != "ABSTRACT":
        headings_sections.append((current_heading.strip(), '\n'.join(current_section)))
    
    return headings_sections

''' 
Input: the file path to a docx protocol file
Output: a dictionary organized where each key is a section of the abstract 
and the value is the text inside that heading
If no abstract is found, the function returns None
'''
def process_abstract(file_path):

    doc = docx.Document(file_path)
    is_abstract = False
    output = {}
    for paragraph in doc.paragraphs:
        #if is a heading and not abstract heading, then quit the loop
        if is_heading(paragraph) and paragraph.text.upper() == "ABSTRACT":
                is_abstract = True
        # this means that we have reached the header after the abstract
        elif is_heading(paragraph) and is_abstract:
            return output
        else:
            # if doesn't contain one of the key words, then the paragraph isn't in the abstract
            # unless there is an extra indent, in which case that should be checked
            # 'break' ensures that each section is only added once
            # Note: does not cover variables section because summary would likely not be useful
            for abs in ABS_HEADINGS:
                if abs in paragraph.text.upper():
                    # separate the title and the text itself
                    chunk = paragraph.text.lower()
                    chunk = chunk.replace(abs.lower() + ':', '')
                    output[abs.lower()] = chunk
                    break
    return None
        

'''
This function should check whether the processing worked properly per document 
If not proper processing, then print out the article name so that it can be evaluated / manually entered
'''
def check_processing():
    print("TO DO")

'''
Output: a path to the zip file after creating a zip file from the json_data folder
Also deletes the json_data folder from the local computer
'''
def create_zip():
    shutil.make_archive("json_data", format='zip', base_dir="json_data")
    shutil.rmtree("json_data\\")
    return os.getcwd() + "\\json_data.zip"

def delete_zip():
    os.remove("json_data.zip")

'''
Input: a list of dictionaries, where each dictionary will become a json file
Output: a list of json file names. A folder will be created (or added to) that contains
the zip files in them. These are intended to be deleted upon user download in the application.'''
def create_jsons(data):
    folder_name = "json_data\\"
    currdir = os.getcwd()
    if not os.path.exists(currdir + "\\" + folder_name):
        os.mkdir(currdir + "\\" + folder_name)
    json_names = []
    currpath = os.path.join(currdir, folder_name)
    # create the json files and keep track of names in a list for future reference
    if type(data) == dict:
        name = "output0.json"
        with open(currpath + name, "w") as output:
            json.dump(data, output)
        json_names.append(name)
    elif type(data) == list:
        for i, d in enumerate(data):
            name = "output" + str(i) + ".json"
            with open(currpath + name, "w") as output:
                json.dump(d, output)
            json_names.append(name)
    return json_names
    



'''
NEXT: INTEGRATE INTO STREAMLIT APP, AND MAKE OUTPUT USING PEGASUS FOR SUMMARY
NEXT: MAKE SEPARATE STREAMLIT APP (OR SOMETHING) TO AUTOMATE MAKING DATASET FOR FINE TUNING
'''
# try putting google colab in here later
def main():
    file_path = "A3921391_Non_Interventional Study_Protocol Amendment 1 (clean)_24NOV2020 GDMS ARTIS.docx"
    output_dict = process_full_word_doc(file_path)
    dir = os.getcwd()
    full_path = os.path.join(dir, 'test.json')
    with open(full_path, "w") as json_file:
        json.dump(output_dict, json_file)